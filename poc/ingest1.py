"""
DDA INGEST AGENT
-----------------
Reads every file in data/raw/ and extracts text content.
Supports: PDF, DOCX, CSV
Saves normalized output to Firestore as dda_artifacts.

Run: python agents/ingest.py
"""

import os
import sys
import json
import hashlib
import csv
from datetime import datetime
from io import StringIO

# Add parent directory to path so we can import gcp_client
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from gcp_client import (
    list_raw_files, save_artifact, get_firestore,
    COL_ARTIFACTS, PROJECT_ID
)

# ── Text Extractors ───────────────────────────────────────────────────────

def extract_pdf(file_path):
    """Extracts text from PDF files."""
    try:
        import pypdf
        text = ""
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text.strip(), len(reader.pages)
    except Exception as e:
        return f"PDF extraction failed: {e}", 0


def extract_docx(file_path):
    """Extracts text from Word documents."""
    try:
        from docx import Document
        doc   = Document(file_path)
        lines = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n".join(lines), len(doc.paragraphs)
    except Exception as e:
        return f"DOCX extraction failed: {e}", 0


def extract_csv(file_path):
    """Extracts text from CSV files — converts rows to readable sentences."""
    try:
        rows = []
        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                rows.append(row)

        if not rows:
            return "Empty CSV file", 0

        # Convert CSV rows to readable text for Gemini to process
        headers = list(rows[0].keys())
        lines   = [f"Columns: {', '.join(headers)}"]
        for row in rows:
            line = " | ".join([f"{k}: {v}" for k, v in row.items()])
            lines.append(line)

        return "\n".join(lines), len(rows)
    except Exception as e:
        return f"CSV extraction failed: {e}", 0


def compute_hash(file_path):
    """Computes MD5 hash to detect duplicate files."""
    with open(file_path, "rb") as f:
        return hashlib.md5(f.read()).hexdigest()


def is_duplicate(content_hash):
    """Checks Firestore if this file was already processed."""
    db   = get_firestore()
    docs = db.collection(COL_ARTIFACTS)\
             .where("content_hash", "==", content_hash)\
             .limit(1)\
             .stream()
    return any(True for _ in docs)


def detect_document_type(filename):
    """Detects document category from filename."""
    name = filename.lower()
    if "contract" in name:
        return "pricing_contract"
    elif "email" in name:
        return "email_thread"
    elif "price_list" in name or "price" in name:
        return "price_list"
    else:
        return "unknown"


# ── Main INGEST Function ──────────────────────────────────────────────────

def ingest_file(file_info):
    """
    Processes a single file end-to-end:
    1. Check for duplicates
    2. Extract text based on file type
    3. Save to Firestore
    """
    filename  = file_info["filename"]
    full_path = file_info["full_path"]
    extension = file_info["extension"]

    print(f"\n  📄 Processing: {filename}")

    # Step 1: Duplicate check
    content_hash = compute_hash(full_path)
    if is_duplicate(content_hash):
        print(f"     ⏭️  Skipping — already processed")
        return None

    # Step 2: Extract text based on file type
    if extension == "pdf":
        raw_text, page_count = extract_pdf(full_path)
        extractor = "pypdf"
    elif extension == "docx":
        raw_text, page_count = extract_docx(full_path)
        extractor = "python-docx"
    elif extension == "csv":
        raw_text, page_count = extract_csv(full_path)
        extractor = "csv-parser"
    else:
        print(f"     ⚠️  Unsupported format: {extension} — skipping")
        return None

    # Step 3: Score extraction quality
    word_count  = len(raw_text.split())
    confidence  = min(1.0, word_count / 50)  # low word count = low confidence
    has_error   = "extraction failed" in raw_text.lower()

    if has_error:
        confidence = 0.0

    # Step 4: Build normalized artifact
    artifact_id = f"{filename.replace('.', '_')}_{content_hash[:8]}"
    artifact    = {
        "artifact_id":       artifact_id,
        "filename":          filename,
        "file_path":         full_path,
        "extension":         extension,
        "document_type":     detect_document_type(filename),
        "raw_text":          raw_text,
        "word_count":        word_count,
        "page_count":        page_count,
        "content_hash":      content_hash,
        "extractor_used":    extractor,
        "extraction_confidence": round(confidence, 2),
        "requires_review":   confidence < 0.5 or has_error,
        "status":            "EXTRACTED" if not has_error else "FAILED",
        "ingested_at":       datetime.utcnow().isoformat(),
        "pipeline_stage":    "INGESTED"
    }

    # Step 5: Save to Firestore
    save_artifact(artifact_id, artifact)

    # Step 6: Report
    status = "✅" if not has_error else "❌"
    print(f"     {status} Extracted {word_count} words | confidence: {confidence:.0%} | saved as: {artifact_id}")

    return artifact


def run_ingest():
    """Main entry point — processes all files in data/raw/."""
    print("\n" + "="*55)
    print("  🏛️  DDA INGEST AGENT")
    print("  Reads raw files → Extracts text → Saves to Firestore")
    print("="*55)

    # Get all raw files
    files = list_raw_files()
    print(f"\n  Found {len(files)} files to process...")

    results = {
        "processed": [],
        "skipped":   [],
        "failed":    []
    }

    for file_info in files:
        artifact = ingest_file(file_info)

        if artifact is None:
            results["skipped"].append(file_info["filename"])
        elif artifact["status"] == "FAILED":
            results["failed"].append(file_info["filename"])
        else:
            results["processed"].append(file_info["filename"])

    # Summary
    print("\n" + "="*55)
    print("  INGEST COMPLETE")
    print("="*55)
    print(f"  ✅ Processed:  {len(results['processed'])} files")
    print(f"  ⏭️  Skipped:    {len(results['skipped'])} files (duplicates)")
    print(f"  ❌ Failed:     {len(results['failed'])} files")
    print(f"\n  📦 All artifacts saved to Firestore")
    print(f"     Collection: dda_artifacts")
    print(f"     Project:    {PROJECT_ID}")

    if results["processed"]:
        print(f"\n  Files processed:")
        for f in results["processed"]:
            print(f"    📄 {f}")

    print("\n  Next step: python agents/correlate.py")
    print("="*55 + "\n")

    return results


if __name__ == "__main__":
    # Install required packages if missing
    try:
        import pypdf
    except ImportError:
        print("Installing pypdf...")
        os.system("pip install pypdf -q")
        import pypdf

    try:
        from docx import Document
    except ImportError:
        print("Installing python-docx...")
        os.system("pip install python-docx -q")

    run_ingest()