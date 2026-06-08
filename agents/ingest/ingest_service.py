"""
DDA INGEST SERVICE — Cloud Run Entry Point
Wraps agents/ingest.py logic as a FastAPI HTTP service.
Triggered by Cloud Pub/Sub push subscription (GCS object notifications).
"""

import os
import json
import base64
import hashlib
import logging
from datetime import datetime, timezone
from fastapi import FastAPI, Request, HTTPException
from google.cloud import firestore, pubsub_v1, storage
import vertexai
from vertexai.generative_models import GenerativeModel
import uvicorn
from docx import Document
import io

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("dda-ingest")

app = FastAPI(title="DDA Ingest Service")

PROJECT_ID      = os.environ["GCP_PROJECT_ID"]      # the-orchestrators
REGION          = os.environ.get("GCP_REGION", "us-central1")
ARTIFACTS_COL   = os.environ.get("FIRESTORE_ARTIFACTS", "dda_artifacts")
PUBSUB_INGEST_LISTEN_TOPIC    = os.environ.get("PUBSUB_INGEST_LISTEN_TOPIC", "dda-ingest-trigger")
PUBSUB_INGEST_TO_CORRELATE_TOPIC = os.environ.get("PUBSUB_INGEST_TO_CORRELATE_TOPIC", "dda-correlate-trigger")
DOC_AI_PROCESSOR = os.environ.get("DOCUMENT_AI_PROCESSOR_ID", "")

db        = firestore.Client(project=PROJECT_ID)
publisher = pubsub_v1.PublisherClient()

# ── helpers ────────────────────────────────────────────────────────────────

def md5_of_gcs(bucket_name: str, blob_name: str) -> str:
    client = storage.Client(project=PROJECT_ID)
    blob   = client.bucket(bucket_name).blob(blob_name)
    blob.reload()
    return blob.md5_hash or hashlib.md5(blob_name.encode()).hexdigest()

def extract_text_from_gcs(bucket_name: str, blob_name: str, mime_type: str) -> dict:
    """
    Route to correct extractor based on mime type.
    Uses Document AI for PDFs/DOCX, raw read for CSV.
    Returns { raw_text, extraction_confidence, extractor_used }
    """
    gcs_uri = f"gs://{bucket_name}/{blob_name}"

    if mime_type == "text/csv" or blob_name.endswith(".csv"):
        client = storage.Client(project=PROJECT_ID)
        data   = client.bucket(bucket_name).blob(blob_name).download_as_text()
        return {"raw_text": data, "extraction_confidence": 1.0, "extractor_used": "csv_reader"}

    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or blob_name.endswith(".docx"):
        client = storage.Client(project=PROJECT_ID)
        content = client.bucket(bucket_name).blob(blob_name).download_as_bytes()
        document = Document(io.BytesIO(content))
        full_text = []
        for para in document.paragraphs:
            full_text.append(para.text)
        return {"raw_text": "\n".join(full_text), "extraction_confidence": 0.95, "extractor_used": "python-docx"}

    if DOC_AI_PROCESSOR:
        from google.cloud import documentai
        dai_client = documentai.DocumentProcessorServiceClient()
        client     = storage.Client(project=PROJECT_ID)
        content    = client.bucket(bucket_name).blob(blob_name).download_as_bytes()
        raw_doc    = documentai.RawDocument(content=content, mime_type=mime_type)
        request    = documentai.ProcessRequest(name=DOC_AI_PROCESSOR, raw_document=raw_doc)
        result     = dai_client.process_document(request=request)
        return {
            "raw_text": result.document.text,
            "extraction_confidence": 0.97,
            "extractor_used": "document_ai"
        }

    # Fallback: plain text read
    client = storage.Client(project=PROJECT_ID)
    data   = client.bucket(bucket_name).blob(blob_name).download_as_text(errors="replace")
    return {"raw_text": data, "extraction_confidence": 0.75, "extractor_used": "raw_read"}

def infer_doc_type(filename: str) -> str:
    fn = filename.lower()
    if "contract" in fn: return "pricing_contract"
    if "email" in fn or "thread" in fn: return "email_thread"
    if "price" in fn or "list" in fn: return "price_list"
    return "unknown"

# ── routes ─────────────────────────────────────────────────────────────────

@app.get("/health")
def health():
    return {"status": "ok", "service": "dda-ingest"}

@app.post("/ingest/trigger")
async def ingest_trigger(request: Request):
    """
    Receives Pub/Sub push message containing artifact_id to ingest.
    """
    body = await request.json()
    logger.info(f"Received Pub/Sub message body: {body}")
    try:
        msg_data = base64.b64decode(body["message"]["data"]).decode()
        logger.info(f"Decoded message data: {msg_data}")
        event = json.loads(msg_data)
        logger.info(f"Parsed event: {event}")
        artifact_id = event["artifact_id"]
    except Exception as e:
        logger.error(f"Bad Pub/Sub message: {e}")
        raise HTTPException(status_code=400, detail=str(e))

    logger.info(f"Processing artifact_id: {artifact_id}")

    artifact_ref = db.collection(ARTIFACTS_COL).document(artifact_id)
    artifact_doc = artifact_ref.get()

    if not artifact_doc.exists:
        logger.error(f"Artifact {artifact_id} not found in Firestore.")
        raise HTTPException(status_code=404, detail=f"Artifact {artifact_id} not found.")

    artifact = artifact_doc.to_dict()

    if artifact.get("pipeline_stage") != "UPLOADED":
        logger.warning(f"Artifact {artifact_id} is not in 'UPLOADED' stage. Current stage: {artifact.get('pipeline_stage')}. Skipping ingestion.")
        return {"status": "skipped", "artifact_id": artifact_id, "reason": "not_in_uploaded_stage"}

    try:
        bucket_name = artifact["gcs_uri"].split("/")[2]
        blob_name   = "/".join(artifact["gcs_uri"].split("/")[3:])
        mime_type   = artifact["mime_type"]

        extraction  = extract_text_from_gcs(bucket_name, blob_name, mime_type)

        updated_fields = {
            "raw_text":               extraction["raw_text"],
            "word_count":             len(extraction["raw_text"].split()),
            "extraction_confidence":  extraction["extraction_confidence"],
            "extractor_used":         extraction["extractor_used"],
            "status":                 "EXTRACTED",
            "pipeline_stage":         "INGESTED",
            "requires_review":        extraction["extraction_confidence"] < 0.65,
            "ingested_at":            datetime.now(timezone.utc).isoformat(),
        }
        artifact_ref.update(updated_fields)
        logger.info(f"Saved artifact: {artifact_id} with INGESTED stage")

        publisher.publish(PUBSUB_INGEST_TO_CORRELATE_TOPIC, json.dumps({"artifact_id": artifact_id}).encode())
        logger.info(f"Published {artifact_id} to {PUBSUB_INGEST_TO_CORRELATE_TOPIC}")

        return {"status": "ok", "artifact_id": artifact_id}

    except Exception as e:
        logger.error(f"Failed to ingest artifact {artifact_id}: {e}", exc_info=True)
        artifact_ref.update({
            "ingest_error": str(e),
            "requires_review": True,
            "pipeline_stage": "INGEST_FAILED",
            "failed_at": datetime.now(timezone.utc).isoformat()
        })
        raise HTTPException(status_code=500, detail=f"Ingestion failed for {artifact_id}: {str(e)}")





if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", 8080)))