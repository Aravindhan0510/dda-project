import os, csv, io, hashlib, json, uuid
from datetime import datetime, timezone
from pathlib import Path
from dotenv import load_dotenv

import pypdf
import docx
from google.cloud import firestore, pubsub_v1, storage
from google.cloud import documentai_v1 as documentai

load_dotenv()

# ── Config ────────────────────────────────────────────────────────────────────
PROJECT_ID          = os.getenv("GCP_PROJECT_ID", "the-orchestrators")
REGION              = os.getenv("GCP_REGION", "us-central1")
FIRESTORE_ARTIFACTS = os.getenv("FIRESTORE_ARTIFACTS", "dda_artifacts")
PUBSUB_INGEST_TOPIC = os.getenv("PUBSUB_INGEST_TOPIC", "dda-ingest-complete")
DOCAI_PROCESSOR_ID  = os.getenv("DOCUMENT_AI_PROCESSOR_ID")
MIN_CONFIDENCE      = float(os.getenv("MIN_EXTRACTION_CONFIDENCE", "0.65"))
GCS_STAGING_BUCKET  = os.getenv("GCS_STAGING_BUCKET", "dda-raw-staging")
GCS_PROCESSED_BUCKET= os.getenv("GCS_PROCESSED_BUCKET", "dda-processed")

# ── GCP Clients ───────────────────────────────────────────────────────────────
db        = firestore.Client(project=PROJECT_ID)
publisher = pubsub_v1.PublisherClient()
topic_path= publisher.topic_path(PROJECT_ID, PUBSUB_INGEST_TOPIC)
storage_client = storage.Client(project=PROJECT_ID)

# ── Scan Detection ────────────────────────────────────────────────────────────
def is_scanned_pdf(pdf_bytes: bytes) -> bool:
    try:
        reader      = pypdf.PdfReader(io.BytesIO(pdf_bytes))
        sample      = min(len(reader.pages), 3)
        total_text  = "".join(reader.pages[i].extract_text() or "" for i in range(sample))
        avg_chars   = len(total_text.strip()) / sample if sample else 0
        return avg_chars < 50
    except Exception:
        return True  # if pypdf fails, treat as scanned → Document AI

# ── Extractors ────────────────────────────────────────────────────────────────
def extract_digital_pdf(pdf_bytes: bytes) -> dict:
    reader    = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    pages     = [reader.pages[i].extract_text() or "" for i in range(len(reader.pages))]
    full_text = "\n".join(pages)
    confidence= min(1.0, len(full_text.strip()) / max(len(reader.pages) * 200, 1))
    return {
        "raw_text":             full_text,
        "page_count":           len(reader.pages),
        "extraction_confidence": round(confidence, 3),
        "extractor_used":       "pypdf"
    }

def extract_docx(file_bytes: bytes) -> dict:
    doc   = docx.Document(io.BytesIO(file_bytes))
    text  = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return {
        "raw_text":             text,
        "page_count":           None,
        "extraction_confidence": 1.0 if len(text) > 50 else 0.5,
        "extractor_used":       "python-docx"
    }

def extract_csv(file_bytes: bytes) -> dict:
    content = file_bytes.decode("utf-8", errors="replace")
    rows    = list(csv.DictReader(io.StringIO(content)))
    text    = "\n".join(str(r) for r in rows)
    return {
        "raw_text":             text,
        "page_count":           None,
        "extraction_confidence": 1.0,
        "extractor_used":       "csv_parser",
        "row_count":            len(rows)
    }

def extract_via_document_ai(file_bytes: bytes, mime_type: str) -> dict:
    client  = documentai.DocumentProcessorServiceClient(
        client_options={"api_endpoint": "us-documentai.googleapis.com"}
    )
    raw_doc = documentai.RawDocument(content=file_bytes, mime_type=mime_type)
    request = documentai.ProcessRequest(name=DOCAI_PROCESSOR_ID, raw_document=raw_doc)
    result  = client.process_document(request=request)
    doc     = result.document
    confidence = (
        sum(p.layout.confidence for p in doc.pages) / len(doc.pages)
        if doc.pages else 0.5
    )
    return {
        "raw_text":             doc.text,
        "page_count":           len(doc.pages),
        "extraction_confidence": round(confidence, 3),
        "extractor_used":       "document_ai_ocr"
    }

# ── MIME Router ───────────────────────────────────────────────────────────────
def route_and_extract(file_bytes: bytes, filename: str, mime_type: str) -> dict:
    ext = Path(filename).suffix.lower()

    if ext == ".pdf" or mime_type == "application/pdf":
        if is_scanned_pdf(file_bytes):
            print(f"    → Scanned PDF → Document AI")
            return extract_via_document_ai(file_bytes, "application/pdf")
        else:
            print(f"    → Digital PDF → pypdf")
            return extract_digital_pdf(file_bytes)

    elif ext == ".docx":
        print(f"    → DOCX → python-docx")
        return extract_docx(file_bytes)

    elif ext == ".csv":
        print(f"    → CSV → csv_parser")
        return extract_csv(file_bytes)

    elif ext in (".jpg", ".jpeg", ".png", ".tiff", ".tif"):
        print(f"    → Image → Document AI")
        return extract_via_document_ai(file_bytes, f"image/{ext.strip('.')}")

    else:
        raise ValueError(f"Unsupported file type: {ext}")

# ── Helpers ───────────────────────────────────────────────────────────────────
def compute_md5(data: bytes) -> str:
    return hashlib.md5(data).hexdigest()

def is_duplicate(content_hash: str) -> bool:
    docs = db.collection(FIRESTORE_ARTIFACTS)\
              .where("content_hash", "==", content_hash)\
              .limit(1).get()
    return len(docs) > 0

def infer_document_type(filename: str) -> str:
    fn = filename.lower()
    if "contract" in fn:  return "pricing_contract"
    if "email"    in fn:  return "email_thread"
    if "price"    in fn:  return "price_list"
    if "invoice"  in fn:  return "invoice"
    return "unknown"

def build_artifact_id(filename: str, content_hash: str) -> str:
    stem = Path(filename).stem.replace(" ", "_")
    return f"{stem}_{content_hash[:8]}"

def save_artifact(artifact: dict):
    db.collection(FIRESTORE_ARTIFACTS)\
      .document(artifact["artifact_id"])\
      .set(artifact)

def publish_completion(artifact_id: str):
    payload = json.dumps({
        "artifact_id":    artifact_id,
        "pipeline_stage": "INGESTED",
        "timestamp":      datetime.now(timezone.utc).isoformat()
    }).encode()
    publisher.publish(topic_path, payload)

# ── Main ──────────────────────────────────────────────────────────────────────
def ingest_all():
    bucket = storage_client.bucket(GCS_STAGING_BUCKET)
    blobs   = list(bucket.list_blobs())

    print(f"\n{'='*60}")
    print(f"  DDA INGEST AGENT — Enterprise Build (GCS)")
    print(f"  Processing {len(blobs)} files from gs://{GCS_STAGING_BUCKET}")
    print(f"{'='*60}\n")

    results = {"success": [], "duplicate": [], "failed": [], "review_required": []}

    for blob in blobs:
        filename = blob.name
        print(f"[{filename}]")

        try:
            file_bytes   = blob.download_as_bytes()
            content_hash = compute_md5(file_bytes)

            if is_duplicate(content_hash):
                print(f"    → DUPLICATE — skipping\n")
                results["duplicate"].append(filename)
                continue

            ext_map   = {".pdf": "application/pdf",
                         ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         ".csv":  "text/csv"}
            mime_type = ext_map.get(Path(filename).suffix.lower(), "application/octet-stream")

            extraction = route_and_extract(file_bytes, filename, mime_type)

            artifact_id     = build_artifact_id(filename, content_hash)
            requires_review = extraction["extraction_confidence"] < MIN_CONFIDENCE

            artifact = {
                "artifact_id":            artifact_id,
                "filename":               filename,
                "gcs_uri":                f"gs://{GCS_PROCESSED_BUCKET}/{filename}",
                "document_type":          infer_document_type(filename),
                "raw_text":               extraction["raw_text"],
                "page_count":             extraction.get("page_count"),
                "word_count":             len(extraction["raw_text"].split()),
                "extraction_confidence":  extraction["extraction_confidence"],
                "extractor_used":         extraction["extractor_used"],
                "content_hash":           content_hash,
                "mime_type":              mime_type,
                "status":                 "EXTRACTED",
                "pipeline_stage":         "INGESTED",
                "requires_review":        requires_review,
                "ingested_at":            datetime.now(timezone.utc).isoformat()
            }

            save_artifact(artifact)
            
            # Move to processed bucket
            processed_bucket = storage_client.bucket(GCS_PROCESSED_BUCKET)
            bucket.copy_blob(blob, processed_bucket, new_name=filename)
            blob.delete()
            
            publish_completion(artifact_id)

            status = "⚠  LOW CONFIDENCE — flagged" if requires_review else "✅"
            print(f"    → {status} | extractor: {extraction['extractor_used']} | confidence: {extraction['extraction_confidence']:.2f} | words: {artifact['word_count']}\n")

            if requires_review:
                results["review_required"].append(filename)
            else:
                results["success"].append(filename)

        except Exception as e:
            print(f"    → ❌ FAILED: {e}\n")
            results["failed"].append(filename)

    print(f"{'='*60}")
    print(f"  INGEST COMPLETE")
    print(f"  ✅ Success:          {len(results['success'])}")
    print(f"  ⚠  Review Required:  {len(results['review_required'])}")
    print(f"  🔁 Duplicates:       {len(results['duplicate'])}")
    print(f"  ❌ Failed:           {len(results['failed'])}")
    print(f"{'='*60}\n")
    return results

if __name__ == "__main__":
    ingest_all()
